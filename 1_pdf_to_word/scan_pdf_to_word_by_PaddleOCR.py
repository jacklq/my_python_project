import math
import time
import os
from collections import Counter
import docx  # 注意需安装python-docx这个包
import yaml
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import shared
from docx.oxml.ns import qn

from PIL import Image
from PIL import ImageFile  # 注意需安装pillow这个包
import fitz  # pdf转为图片,注意需安装PyMuPDF这个包

ImageFile.LOAD_TRUNCATED_IMAGES = True
Image.MAX_IMAGE_PIXELS = None

import cv2  # 注意需安装opencv-python这个包

from paddleocr import PPStructure, save_structure_res
from paddleocr.ppstructure.recovery.recovery_to_doc import sorted_layout_boxes
from paddleocr.ppstructure.recovery.table_process import HtmlToDocx

"""

处理流程：
    将扫描版按照每页转换成图片，读取每个图片采用ppstructure算法将图片转换为word
参考文档：
paddleocr地址：https://github.com/PaddlePaddle/PaddleOCR/blob/release/2.6/README_ch.md
ppstructure地址：https://github.com/PaddlePaddle/PaddleOCR/tree/release/2.6/ppstructure
            快速开始教程：https://github.com/PaddlePaddle/PaddleOCR/blob/release/2.6/ppstructure/docs/quickstart.md
"""
save_folder = './output'
'''
将PDF转化为图片，
pdfPath pdf文件的路径，
imgPath 图像要保存的路径，
zoom_x x方向的缩放系数
zoom_y y方向的缩放系数
rotation_angle 旋转角度
zoom_x和zoom_y一般取相同值，值越大，图像分辨率越高
返回目标pdf的名称和页数，便于下一步操作
'''


def convert_pdf_to_image(zoom_x=10, zoom_y=10, rotation_angle=0):
    start = time.time()
    # 获取pdf文件名称
    name = pdf_path.split("/")[-1].split('.')[0]
    # 打开PDF文件
    pdf = fitz.open(pdf_path)
    # 获取pdf页数
    num = pdf.page_count
    # 逐页读取PDF
    for pg in range(0, num):
        page = pdf[pg]
        # 设置缩放和旋转系数
        trans = fitz.Matrix(zoom_x, zoom_y).prerotate(rotation_angle)
        pm = page.get_pixmap(matrix=trans, alpha=False)
        # 开始写图像
        pm.save(word_file_path + name + "_" + str(pg) + ".png")

    pdf.close()
    end = time.time()
    print('convert_pdf_to_image Running time: %s Seconds' % (end - start))
    return name, num


# 计算两个颜色之间的欧几里得距离
def color_distance(c1, c2):
    r1, g1, b1 = c1
    r2, g2, b2 = c2
    return math.sqrt((r1 - r2) ** 2 + (g1 - g2) ** 2 + (b1 - b2) ** 2)


# 转化图像为白底黑字（方式2：调用convert函数转换成黑白色，会将彩色的变为黑白色，并不会删除，推荐使用）
def convert_image_to_black(img_str):
    start = time.time()
    new_img_path = img_str.split(".")[0] + "_new.png"

    # 打开原图
    img = Image.open(img_str)
    # 将图片转成灰度模式即黑白色
    im_gray = img.convert("RGB")
    # 保存新图像
    im_gray.save(new_img_path)
    end = time.time()
    print('convert_image_to_black Running time: %s Seconds' % (end - start))
    return new_img_path


def convert_info_docx(res_all, save_folder, img_name, doc):
    for i, each_res in enumerate(res_all):
        img_idx = each_res['img_idx']

        if each_res['type'].lower() == 'figure':
            excel_save_folder = os.path.join(save_folder, img_name)
            img_path = os.path.join(excel_save_folder, '{}_{}.jpg'.format(each_res['bbox'], img_idx))
            paragraph_pic = doc.add_paragraph()
            paragraph_pic.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph_pic.add_run("")
            run.add_picture(img_path, width=shared.Inches(5))

        elif each_res['type'].lower() == 'title':
            title_all = ""
            for i in range(len(each_res['res'])):
                title_all += each_res['res'][i]['text']
            set_title_formate(doc, title_all)
        elif each_res['type'].lower() == 'table':
            parser = HtmlToDocx()
            parser.table_style = 'TableGrid'
            parser.handle_table(each_res['res']['html'], doc)
        else:
            # 判断是不是新段标志
            is_new_para = get_is_new_para_flag(each_res)
            # 将每行合并成每个段落
            new_para_all = combine_row_to_paragraph(is_new_para, each_res['res'])
            # 设置正文格式
            set_text_formate(doc, new_para_all, is_new_para)


# 通过文字的起始位置判断是不是新段的开始，true代表是新段的第一行
def get_is_new_para_flag(text_all):
    text_region_first_x_num = []
    is_new_para = []

    for i in range(len(text_all['res'])):
        text_region_first_x_num.append(text_all['res'][i]['text_region'][0][0])
    min_num = min(text_region_first_x_num)
    for i in range(len(text_region_first_x_num)):
        if text_region_first_x_num[i] < min_num + 200:  # 只要前面有两个空格，就认为是新段落，这的200是个大约数
            is_new_para.append(False)  # false代表非新段落
        else:
            is_new_para.append(True)
    return is_new_para


def image_to_word():
    name, num = convert_pdf_to_image(zoom_x=10, zoom_y=10, rotation_angle=0)
    doc = docx.Document()
    table_engine = PPStructure(recovery=True, lang='ch')
    for pg in range(0, num):
        each_img_path = word_file_path + name + "_" + str(pg) + ".png"
        # 将图片转成黑白色
        new_img_path = convert_image_to_black(each_img_path)
        # 调用ppstructure将图片转成文字
        each_img = cv2.imread(new_img_path)
        result = table_engine(each_img)
        save_structure_res(result, save_folder, os.path.basename(new_img_path).split('.')[0])
        for line in result:
            line.pop('img')
            print(line)
        h, w, _ = each_img.shape
        res = sorted_layout_boxes(result, w)
        convert_info_docx(res, save_folder, os.path.basename(new_img_path).split('.')[0], doc)
        # 一页结束后新增分页符
        doc.add_section()
    set_section_format(doc)
    doc.save(word_file_path + name + ".docx")


# 设置正文格式,
def set_text_formate(doc, new_para_all, is_new_para):
    for j, each_new_para in enumerate(new_para_all):
        text_para = doc.add_paragraph()
        text_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
        text_para.paragraph_format.line_spacing = Pt(30)  # 行间距，固定值30磅
        if j == 0 and is_new_para[0] == False:  # 第一行是没空格，说明是顶格，非新段第一行，所以不空格
            text_para.paragraph_format.first_line_indent = Pt(0)
        else:
            text_para.paragraph_format.first_line_indent = Pt(32)  # 首行缩进32磅即空两格
        text_para.paragraph_format.space_before = Pt(0)  # 段前0磅
        text_para.paragraph_format.space_after = Pt(0)  # 段后0磅
        text_formate = text_para.add_run(each_new_para)
        text_formate.font.size = Pt(16)  # 设置字体大小三号
        text_formate.bold = False  # 设置字体是否加粗
        text_formate.font.name = 'Times New Roman'  # 设置西文字体
        text_formate.element.rPr.rFonts.set(qn('w:eastAsia'), 'GB2312')  # 设置中文字体


def combine_row_to_paragraph(is_new_para, text_all):
    new_para_all = []
    each_new_para_temp = ""
    for i, each_text in enumerate(text_all):
        if is_new_para[i]:
            if each_new_para_temp != "":
                new_para_all.append(each_new_para_temp)
                each_new_para_temp = ""
            each_new_para_temp = each_text['text']
        else:
            each_new_para_temp = each_new_para_temp + each_text['text']
    new_para_all.append(each_new_para_temp)
    return new_para_all


# 设置标题格式
def set_title_formate(doc, title):
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    title_para.paragraph_format.line_spacing = Pt(30)  # 行间距，固定值30磅
    title_text = title_para.add_run(title)
    title_text.font.size = Pt(22)  # 设置字体大小
    title_text.bold = False  # 设置字体是否加粗
    title_text.font.name = 'Times New Roman'  # 设置西文字体
    title_text.element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')  # 设置中文字体


# 设置页面布局
def set_section_format(doc):
    for sec in doc.sections:
        # 设置页面边距
        sec.top_margin = Cm(3.4)
        sec.left_margin = Cm(2.8)
        sec.right_margin = Cm(2.6)
        sec.bottom_margin = Cm(3.2)
        # sec.top_margin = Cm(1.7)
        # sec.left_margin = Cm(1.8)
        # sec.right_margin = Cm(1.6)
        # sec.bottom_margin = Cm(1.5)
        # 设置纸张大小(A4)
        sec.page_height = Mm(297)
        sec.page_width = Mm(210)
        # 设置页眉页脚距离
        sec.header_distance = Cm(1.5)
        sec.footer_distance = Cm(0.2)


if __name__ == "__main__":
    # pdf_path = input("请输入pdf存储路径及名称(例如C:/Users/jack8/Desktop/专业证明.pdf)： ")
    # word_file_path = input("请输入生成word存储地址(例如C:/Users/jack8/Desktop/)： ")
    with open('config.yaml', 'r') as f:  # 用with读取文件更好
        configs = yaml.load(f, Loader=yaml.FullLoader)  # 按字典格式读取并返回
        # 显示读取后的内容
    pdf_path = configs["path"]["pdf_path"]
    word_file_path = configs["path"]["word_file_path"]
    image_to_word()
