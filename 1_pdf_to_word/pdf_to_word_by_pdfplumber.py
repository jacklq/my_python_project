import os
import pdfplumber
import docx  # 注意需安装python-docx这个包
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn
import xlwt

"""
 通过pdfplumber提取pdf中的文字（试用于正常的pdf即由word另存为pdf得来的，不能提取扫描件）
 注意：
    1、pdfplumber不能提取扫描件
    2、docx 需安装python-docx这个包
"""

name = ""
result_file_path=""

# 删除正文多余的换行符
def delete_needless_newline_char(replace_pos, text):

    text_final = text[:replace_pos[0]]
    # 如果只有一个换行符，则去除这个换行符即可
    if len(replace_pos) == 1:
        text_final = text_final + text[replace_pos[0] + 1:]
    # 如果有多个换行符
    for i in range(1, len(replace_pos)):#rang函数不包括最后个数，如len函数结果为15，则i最大为14
        # 取两个换行符之间的字符串
        temp_replace_str = text[replace_pos[i - 1] + 1:replace_pos[i]]
        text_final = text_final + temp_replace_str
        #如果是最后一个换行符，需要将最后一个换行符到字符串末尾这一个子字符串加上
        if i==len(replace_pos)-1:
            last_str=text[replace_pos[i]+1:]
            text_final = text_final + last_str
    print(text_final)
    return text_final


# 查找要替换掉和要保留的换行符位置
def find_replace_and_retain_pos(text, str):
    replace_pos = []  # 要替换掉的换行符位置
    retain_pos = []  # 要保留的换行符位置
    for i in range(0, len(text)):
        if text[i] == str and text[i - 1] != "。":  # 是换行符且前边一个字符不是句号，那么替换成空格
            replace_pos.append(i)
        if text[i] == str and text[i - 1] == "。":  # 是换行符且前边一个字符是句号，那么保留
            retain_pos.append(i)

    return replace_pos, retain_pos


# 将提取的文字进行预处理（提取标题和正文，删除正文中多余的换行符）
def text_all_preprocessing(text_all):
    first_num = text_all.find('\n')  # 第一个换行符位置
    title = text_all[:first_num + 1]  # 获取标题，认为第一个换行符前边的就是标题
    text = text_all[first_num + 1:]  # 获取正文
    # 获取换行符位置（包含要替换掉的和要保留的）
    replace_pos, retain_pos = find_replace_and_retain_pos(text, "\n")
    # 删除正文多余的换行符
    text_final = delete_needless_newline_char(replace_pos, text)
    return text_final, title


 # 设置正文格式,通过识别\n分割成多个段落
def set_text_formate(doc, text):
    text_more_para = text.split('\n')
    for i in range(0, len(text_more_para)):
        p1 = doc.add_paragraph()
        p1.paragraph_format.line_spacing = Pt(30)  # 行间距，固定值30磅
        p1.paragraph_format.first_line_indent = Pt(32)  # 首行缩进32磅
        p1.paragraph_format.space_before = Pt(0)  # 段前30磅
        p1.paragraph_format.space_after = Pt(0)  # 段后15磅
        text1 = p1.add_run(text_more_para[i])
        text1.font.size = Pt(16)  # 设置字体大小
        text1.bold = False  # 设置字体是否加粗
        text1.font.name = 'Times New Roman'  # 设置西文字体
        text1.element.rPr.rFonts.set(qn('w:eastAsia'), 'GB2312')  # 设置中文字体

# 设置标题格式
def set_title_formate(doc, title):
    t1 = doc.add_paragraph()
    t1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    t1.paragraph_format.line_spacing = Pt(30)  # 行间距，固定值30磅
    text1 = t1.add_run(title)
    text1.font.size = Pt(22)  # 设置字体大小
    text1.bold = False  # 设置字体是否加粗
    text1.font.name = 'Times New Roman'  # 设置西文字体
    text1.element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')  # 设置中文字体


# 将读取的文字写入到word中
def write_to_word(text_all):
    # 将提取的文字进行预处理（提取标题和正文，删除正文中多余的换行符）
    text_final, title = text_all_preprocessing(text_all)
    doc = docx.Document()
    # 设置标题格式
    set_title_formate(doc, title)
    # 设置正文格式,通过识别\n分割成多个段落
    set_text_formate(doc, text_final)
    doc.save(result_file_path + name + ".docx")

    doc.save(result_file_path + name + ".docx")


# 将读取的文字写入到excel中
def write_to_excel(table_list):
    # 创建Excel表对象
    workbook = xlwt.Workbook(encoding='utf8')
    # 新建sheet表
    worksheet = workbook.add_sheet('Sheet1')
    # 自定义列名
    col1 = table_list[0]

    # 将列属性元组col写进sheet表单中第一行
    for i in range(0, len(col1)):
        worksheet.write(0, i, col1[i])
    # 将数据写进sheet表单中
    for i in range(0, len(table_list[1:])):
        data = table_list[1:][i]
        for j in range(0, len(col1)):
            worksheet.write(i + 1, j, data[j])
    # 存文件分两种格式
    workbook.save(result_file_path + name + ".xls")


# 读取pdf内容为txt，提取表格数据
def read_from_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        # 提取文字
        content = ""
        for page in pdf.pages:
            content = content + page.extract_text()
        if content!="":
            write_to_word(content)

        # 提取表格
        table_list = []
        for page in pdf.pages:
            tables = page.extract_tables()
            if len(tables) == 0:
                continue
            table_list.extend(tables)
        if len(tables) >0:
            write_to_excel(table_list)

        return (content, table_list)


if __name__ == "__main__":
    pdf_path = input("请输入pdf存储路径及名称(例如C:/Users/jack8/Desktop/专业证明.pdf)： ")
    name = pdf_path.split("/")[-1].split(".")[0]
    print(name)
    result_file_path = input("请输入生成word或excel存储地址(例如C:/Users/jack8/Desktop/)： ")
    content, table_list = read_from_pdf(pdf_path)
#C:/Users/jack8/Desktop/2022年度山东公安科学技术进步奖揭晓 -简讯.pdf