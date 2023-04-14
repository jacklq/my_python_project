import math
import time

import docx  # 注意需安装python-docx这个包
from docx.shared import Pt
from docx.oxml.ns import qn

from PIL import Image
import fitz  # pdf转为图片,注意需安装PyMuPDF这个包
import pytesseract

from PIL import ImageFile

ImageFile.LOAD_TRUNCATED_IMAGES = True
Image.MAX_IMAGE_PIXELS = None
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

"""

使用说明：
    用于将扫描的pdf转换为word

"""

word_file_path = ""
pdf_path = ""

'''
将PDF转化为图片
pdfPath pdf文件的路径
imgPath 图像要保存的路径
zoom_x x方向的缩放系数
zoom_y y方向的缩放系数
rotation_angle 旋转角度
zoom_x和zoom_y一般取相同值，值越大，图像分辨率越高
返回目标pdf的名称和页数，便于下一步操作
'''


def convert_pdf_to_image(zoom_x=10, zoom_y=10, rotation_angle=0):
    start = time.time()
    # 获取pdf文件名称
    name = pdf_path.split("/")[-1].split('.')[0]
    # 打开PDF文件
    pdf = fitz.open(pdf_path)
    # 获取pdf页数
    num = pdf.page_count
    # 逐页读取PDF
    for pg in range(0, num):
        page = pdf[pg]
        # 设置缩放和旋转系数
        trans = fitz.Matrix(zoom_x, zoom_y).prerotate(rotation_angle)
        pm = page.get_pixmap(matrix=trans, alpha=False)
        # 开始写图像
        pm.save(word_file_path + name + "_" + str(pg) + ".png")

    pdf.close()
    end = time.time()
    print('convert_pdf_to_image Running time: %s Seconds' % (end - start))
    return name, num


# 计算两个颜色之间的欧几里得距离
def color_distance(c1, c2):
    r1, g1, b1 = c1
    r2, g2, b2 = c2
    return math.sqrt((r1 - r2) ** 2 + (g1 - g2) ** 2 + (b1 - b2) ** 2)


# 转化图像为白底黑字（方式1：遍历像素点赋值黑像素，会将彩色的删除，不推荐使用），一定要转化，能提高识别准确性
def transformedImage(img_str):
    start = time.time()
    new_img_path = img_str.split(".")[0] + "_new.png"

    # 打开原图
    img = Image.open(img_str)
    # 创建一个白色的背景图像
    bg_img = Image.new('RGBA', img.size, (255, 255, 255))
    bg_img.save(img_str.split(".")[0] + "_new1.png")
    # 定义相似颜色的阈值，5~200之间为最佳值，5~500为有效值
    threshold = 100

    # 遍历所有像素点
    for x in range(img.width):
        for y in range(img.height):
            # 获取当前像素点的颜色
            color = img.getpixel((x, y))
            # 如果原图当前坐标颜色与给定颜色相似，则在背景图中相同的坐标写入黑色像素点
            if color_distance(color, (0, 0, 0)) < threshold:
                bg_img.putpixel((x, y), (0, 0, 0))

    # 保存新图像
    bg_img.save(new_img_path)
    end = time.time()
    print('transformedImage Running time: %s Seconds' % (end - start))
    return new_img_path


# 转化图像为白底黑字（方式2：调用convert函数转换成黑白色，会将彩色的变为黑白色，并不会删除，推荐使用）
def convert_image_to_black(img_str):
    start = time.time()
    new_img_path = img_str.split(".")[0] + "_new.png"

    # 打开原图
    img = Image.open(img_str)
    # 将图片转成灰度模式即黑白色
    im_gray = img.convert("RGB")
    # 保存新图像
    im_gray.save(new_img_path)
    end = time.time()
    print('convert_image_to_black Running time: %s Seconds' % (end - start))
    return new_img_path


def image_to_word():
    name, num = convert_pdf_to_image(zoom_x=10, zoom_y=10, rotation_angle=0)
    doc = docx.Document()
    for pg in range(0, num):
        each_img_path = word_file_path + name + "_" + str(pg) + ".png"
        # 将图片转成黑白色
        new_img_path = convert_image_to_black(each_img_path)
        # 调用pytesseract将图片转成文字
        text_all = pytesseract.image_to_string(Image.open(new_img_path), lang="chi_sim+eng+num", config="--psm 6")
        # 将提取的文字进行预处理（提取标题和正文，删除正文中多余的换行符）
        text_final, title = text_all_preprocessing(text_all, pg)
        # 将提取的标题和处理的正文写入word
        write_to_word(doc, title, text_final)
    doc.save(word_file_path + name + ".docx")


# 将提取的文字进行预处理（提取标题和正文，删除正文中多余的换行符）
def text_all_preprocessing(text_all, pg):
    first_num = text_all.find('\n')  # 第一个换行符位置
    # 只需在第一页提取标题，其他页码为正文
    if pg == 0:
        # 获取标题，认为第一个换行符前边的就是标题
        title = text_all[:first_num + 1]
        text = text_all[first_num + 1:]
    else:
        title = ""
        text = text_all

    # 获取换行符位置（包含要替换掉的和要保留的）
    replace_pos, retain_pos = find_replace_and_retain_pos(text, "\n")
    # 删除正文多余的换行符
    text_final = delete_needless_newline_char(replace_pos, text)
    return text_final, title


# 删除正文多余的换行符
def delete_needless_newline_char(replace_pos, text):
    # text_final赋初值
    text_final = text[:replace_pos[0]]
    # 如果只有一个换行符，则去除这个换行符即可
    if len(replace_pos) == 1:
        text_final = text_final + text[replace_pos[0] + 1:]
    # 如果有多个换行符，
    for i in range(1, len(replace_pos)):  # rang函数不包括最后个数，如len函数结果为15，则i最大为14
        # 取两个换行符之间的字符串
        temp_replace_str = text[replace_pos[i - 1] + 1:replace_pos[i]]
        text_final = text_final + temp_replace_str
        # 如果是最后一个换行符，需要将最后一个换行符到字符串末尾这一个子字符串加上
        if i == len(replace_pos) - 1:
            last_str = text[replace_pos[i] + 1:]
            text_final = text_final + last_str
    print(text_final)
    return text_final


# 查找要替换掉和要保留的换行符位置
def find_replace_and_retain_pos(text, str):
    replace_pos = []  # 要替换掉的换行符位置
    retain_pos = []  # 要保留的换行符位置
    for i in range(0, len(text)):
        if text[i] == str and text[i - 1] != "。":  # 是换行符且前边一个字符不是句号，那么替换成空格
            replace_pos.append(i)
        if (text[i] == str and text[i - 1] == "。") or (text[i] == str and text[i - 1] == ":"):  # 是换行符且前边一个字符是句号，那么保留
            retain_pos.append(i)

    return replace_pos, retain_pos


# 将读取的文字写入到word中，设置word格式
def write_to_word(doc, title, text):
    # title非空的时候才会设置标题格式
    if title != "":
        # 设置标题格式
        set_title_formate(doc, title)
    # 设置正文格式,通过识别\n分割成多个段落
    set_text_formate(doc, text)


# 设置正文格式,通过识别\n分割成多个段落
def set_text_formate(doc, text):
    text_more_para = text.split('\n')
    for i in range(0, len(text_more_para)):
        p1 = doc.add_paragraph()
        p1.paragraph_format.line_spacing = Pt(30)  # 行间距，固定值30磅
        p1.paragraph_format.first_line_indent = Pt(32)  # 首行缩进32磅
        p1.paragraph_format.space_before = Pt(0)  # 段前30磅
        p1.paragraph_format.space_after = Pt(0)  # 段后15磅
        text1 = p1.add_run(text_more_para[i])
        text1.font.size = Pt(16)  # 设置字体大小
        text1.bold = False  # 设置字体是否加粗
        text1.font.name = 'Times New Roman'  # 设置西文字体
        text1.element.rPr.rFonts.set(qn('w:eastAsia'), 'GB2312')  # 设置中文字体


# 设置标题格式
def set_title_formate(doc, title):
    t1 = doc.add_paragraph()
    t1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    t1.paragraph_format.line_spacing = Pt(30)  # 行间距，固定值30磅
    text1 = t1.add_run(title)
    text1.font.size = Pt(22)  # 设置字体大小
    text1.bold = False  # 设置字体是否加粗
    text1.font.name = 'Times New Roman'  # 设置西文字体
    text1.element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')  # 设置中文字体


if __name__ == "__main__":
    pdf_path = input("请输入pdf存储路径及名称(例如C:/Users/jack8/Desktop/专业证明.pdf)： ")
    word_file_path = input("请输入生成word存储地址(例如C:/Users/jack8/Desktop/)： ")

    image_to_word()
