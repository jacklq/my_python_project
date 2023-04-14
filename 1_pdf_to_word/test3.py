"""
------------------------------------
# @FileName    :test3.py
# @Time        :2023/4/11 21:47
# @Author      :jack
# @description :
------------------------------------
"""
import os
import cv2

from paddleocr import PPStructure,save_structure_res
from paddleocr.ppstructure.recovery.recovery_to_doc import sorted_layout_boxes
from docx import Document
from docx import shared
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

import logging
def convert_info_docx(img, res_all, save_folder, img_name):
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = shared.Pt(6.5)


    text_region_first_x_num = []
    is_new_para = []
    for i, each_res in enumerate(res_all):
        if each_res['type'].lower() == 'text':
            for i in range(len(each_res['res'])):
                text_region_first_x_num.append(each_res['res'][i]['text_region'][0][0])
    min_num = min(text_region_first_x_num)
    for i in range(len(text_region_first_x_num)):
        if text_region_first_x_num[i] < min_num + 200:
            is_new_para.append(False)
        else:
            is_new_para.append(True)
    flag = 1
    for i, each_res in enumerate(res_all):
        img_idx = each_res['img_idx']
        # if flag == 2 and region['layout'] == 'single':
        #     section = doc.add_section(WD_SECTION.CONTINUOUS)
        #     section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '1')
        #     flag = 1
        # elif flag == 1 and region['layout'] == 'double':
        #     section = doc.add_section(WD_SECTION.CONTINUOUS)
        #     section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '2')
        #     flag = 2

        if each_res['type'].lower() == 'figure':
            excel_save_folder = os.path.join(save_folder, img_name)
            img_path = os.path.join(excel_save_folder,
                                    '{}_{}.jpg'.format(each_res['bbox'], img_idx))
            paragraph_pic = doc.add_paragraph()
            paragraph_pic.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph_pic.add_run("")
            if flag == 1:
                run.add_picture(img_path, width=shared.Inches(5))
            elif flag == 2:
                run.add_picture(img_path, width=shared.Inches(2))
        elif each_res['type'].lower() == 'title':
            doc.add_heading(each_res['res'][0]['text'])
        # elif region['type'].lower() == 'table':
            # parser = HtmlToDocx()
            # parser.table_style = 'TableGrid'
            # parser.handle_table(region['res']['html'], doc)
        else:
            paragraph = doc.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            for i, line in enumerate(each_res['res']):
                if i == 0:
                    paragraph_format.first_line_indent = shared.Inches(0.25)
                text_run = paragraph.add_run(line['text'] + ' ')
                text_run.font.size = shared.Pt(10)

    # save to docx
    docx_path = os.path.join(save_folder, '{}_ocr.docx'.format(img_name))
    doc.save(docx_path)
    # logger.info('docx save to {}'.format(docx_path))


if __name__ == '__main__':
    # 中文测试图
    table_engine = PPStructure(recovery=True,lang='ch')
    # 英文测试图
    # table_engine = PPStructure(recovery=True, lang='en')

    save_folder = './output'
    img_path = 'C:/Users/jack8/Desktop/asdasda.png'
    img = cv2.imread(img_path)
    result = table_engine(img)
    save_structure_res(result, save_folder, os.path.basename(img_path).split('.')[0])

    for line in result:
        line.pop('img')
        print(line)

    h, w, _ = img.shape
    res = sorted_layout_boxes(result, w)
    convert_info_docx(img, res, save_folder, os.path.basename(img_path).split('.')[0])