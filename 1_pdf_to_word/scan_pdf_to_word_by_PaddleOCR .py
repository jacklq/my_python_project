import math
import time

import docx  # 注意需安装python-docx这个包
from docx.shared import Pt
from docx.oxml.ns import qn

from PIL import Image
import fitz  # pdf转为图片,注意需安装PyMuPDF这个包
import pytesseract

from PIL import ImageFile

ImageFile.LOAD_TRUNCATED_IMAGES = True
Image.MAX_IMAGE_PIXELS = None
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import cv2

from paddleocr import PPStructure, save_structure_res
from paddleocr.ppstructure.recovery.recovery_to_doc import sorted_layout_boxes
from docx import Document
from docx import shared
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from paddleocr.ppstructure.recovery.table_process import HtmlToDocx

"""

使用说明：
    用于将扫描的pdf转换为word

"""

word_file_path = "C:/Users/jack8/Desktop/"
pdf_path = "C:/Users/jack8/Desktop/sadasdadadasdasdasd.pdf"
save_folder = './output'
'''
将PDF转化为图片
pdfPath pdf文件的路径
imgPath 图像要保存的路径
zoom_x x方向的缩放系数
zoom_y y方向的缩放系数
rotation_angle 旋转角度
zoom_x和zoom_y一般取相同值，值越大，图像分辨率越高
返回目标pdf的名称和页数，便于下一步操作
'''


def convert_pdf_to_image(zoom_x=10, zoom_y=10, rotation_angle=0):
    start = time.time()
    # 获取pdf文件名称
    name = pdf_path.split("/")[-1].split('.')[0]
    # 打开PDF文件
    pdf = fitz.open(pdf_path)
    # 获取pdf页数
    num = pdf.page_count
    # 逐页读取PDF
    for pg in range(0, num):
        page = pdf[pg]
        # 设置缩放和旋转系数
        trans = fitz.Matrix(zoom_x, zoom_y).prerotate(rotation_angle)
        pm = page.get_pixmap(matrix=trans, alpha=False)
        # 开始写图像
        pm.save(word_file_path + name + "_" + str(pg) + ".png")

    pdf.close()
    end = time.time()
    print('convert_pdf_to_image Running time: %s Seconds' % (end - start))
    return name, num


# 计算两个颜色之间的欧几里得距离
def color_distance(c1, c2):
    r1, g1, b1 = c1
    r2, g2, b2 = c2
    return math.sqrt((r1 - r2) ** 2 + (g1 - g2) ** 2 + (b1 - b2) ** 2)


# 转化图像为白底黑字（方式2：调用convert函数转换成黑白色，会将彩色的变为黑白色，并不会删除，推荐使用）
def convert_image_to_black(img_str):
    start = time.time()
    new_img_path = img_str.split(".")[0] + "_new.png"

    # 打开原图
    img = Image.open(img_str)
    # 将图片转成灰度模式即黑白色
    im_gray = img.convert("RGB")
    # 保存新图像
    im_gray.save(new_img_path)
    end = time.time()
    print('convert_image_to_black Running time: %s Seconds' % (end - start))
    return new_img_path


# 设置正文格式,通过识别\n分割成多个段落
def set_text_formate(doc, text_all, is_new_para):
    for i, each_text in enumerate(text_all):
        text_para = doc.add_paragraph()
        text_para.paragraph_format.line_spacing = Pt(30)  # 行间距，固定值30磅
        if is_new_para[i]:  # 若是新段第一行则空两格
            text_para.paragraph_format.first_line_indent = Pt(32)  # 首行缩进32磅即空两格
        text_para.paragraph_format.space_before = Pt(0)  # 段前0磅
        text_para.paragraph_format.space_after = Pt(0)  # 段后0磅
        print(each_text['text'])
        text_formate = text_para.add_run(each_text['text'])
        text_formate.font.size = Pt(16)  # 设置字体大小
        text_formate.bold = False  # 设置字体是否加粗
        text_formate.font.name = 'Times New Roman'  # 设置西文字体
        text_formate.element.rPr.rFonts.set(qn('w:eastAsia'), 'GB2312')  # 设置中文字体



def convert_info_docx(img, res_all, save_folder, img_name, doc):
    text_region_first_x_num = []
    is_new_para = []
    for i, each_res in enumerate(res_all):
        if each_res['type'].lower() == 'text':
            for i in range(len(each_res['res'])):
                text_region_first_x_num.append(each_res['res'][i]['text_region'][0][0])
    min_num = min(text_region_first_x_num)
    for i in range(len(text_region_first_x_num)):
        if text_region_first_x_num[i] < min_num + 200:
            is_new_para.append(False)
        else:
            is_new_para.append(True)

    for i, each_res in enumerate(res_all):
        img_idx = each_res['img_idx']

        if each_res['type'].lower() == 'figure':
            excel_save_folder = os.path.join(save_folder, img_name)
            img_path = os.path.join(excel_save_folder, '{}_{}.jpg'.format(each_res['bbox'], img_idx))
            paragraph_pic = doc.add_paragraph()
            paragraph_pic.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph_pic.add_run("")
            run.add_picture(img_path, width=shared.Inches(5))

        elif each_res['type'].lower() == 'title':
            title_all = ""
            for i in range(len(each_res['res'])):
                title_all += each_res['res'][i]['text']
            set_title_formate(doc, title_all)
        elif each_res['type'].lower() == 'table':
            parser = HtmlToDocx()
            parser.table_style = 'TableGrid'
            parser.handle_table(each_res['res']['html'], doc)
        else:
            set_text_formate(doc, each_res['res'], is_new_para)


def image_to_word():
    name, num = convert_pdf_to_image(zoom_x=10, zoom_y=10, rotation_angle=0)
    doc = docx.Document()
    table_engine = PPStructure(recovery=True, lang='ch')
    for pg in range(0, num):
        each_img_path = word_file_path + name + "_" + str(pg) + ".png"
        # 将图片转成黑白色
        new_img_path = convert_image_to_black(each_img_path)
        # 调用pytesseract将图片转成文字
        each_img = cv2.imread(new_img_path)
        result = table_engine(each_img)
        save_structure_res(result, save_folder, os.path.basename(new_img_path).split('.')[0])
        for line in result:
            line.pop('img')
            print(line)

        h, w, _ = each_img.shape
        res = sorted_layout_boxes(result, w)
        convert_info_docx(each_img, res, save_folder, os.path.basename(new_img_path).split('.')[0], doc)

    doc.save(word_file_path + name + ".docx")




# 设置标题格式
def set_title_formate(doc, title):
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    title_para.paragraph_format.line_spacing = Pt(30)  # 行间距，固定值30磅
    title_text = title_para.add_run(title)
    title_text.font.size = Pt(22)  # 设置字体大小
    title_text.bold = False  # 设置字体是否加粗
    title_text.font.name = 'Times New Roman'  # 设置西文字体
    title_text.element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')  # 设置中文字体


if __name__ == "__main__":
    # pdf_path = input("请输入pdf存储路径及名称(例如C:/Users/jack8/Desktop/专业证明.pdf)： ")
    # word_file_path = input("请输入生成word存储地址(例如C:/Users/jack8/Desktop/)： ")

    image_to_word()


#C:/Users/jack8/Desktop/asdasda.pdf