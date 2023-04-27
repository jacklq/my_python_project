
"""
------------------------------------
# @FileName    :crawl_web_data.py
# @Time        :2023/4/17 19:40
# @Author      :jack
# @description :

------------------------------------
"""

"""实现流程说明：
1、给定通知模块网址，在这个网址上各通知信息分为多页
2、获取该页所有具体通知内容的url
3、遍历获取的具体url，获取每个通知的标题、正文、附件
4、分为两种处理模式，由store_text_to_one_word_flag变量决定，true代表汇总到一个word,false代表分开存放：
    模式一：将所有的标题、正文汇总到一个word中
    模式二：将每个通知单独存放，如果有附件（一个或多个），则新建一个文件夹，存放含有标题、正文的word和附件（附件1，附件2.。。）
    将爬取记录存放到新建文件中，每次在遍历通知时，都要判断是否已爬取过，若爬取过则跳过，进行下一条

"""
from docxcompose.composer import Composer
import os
import sys

from urllib.request import urlretrieve
import bs4
from docx.oxml.ns import qn
import docx
import requests
from bs4 import BeautifulSoup
import yaml  # 安装pyyaml
import re  # 导入正则表达式库

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


def set_title_formate(doc, title):
    head = doc.add_heading("", level=1)  # 这里不填标题内容
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 设置为居中
    head_formate = head.add_run(title)
    head_formate.font.name = 'Times New Roman'  # 设置西文字体
    head_formate.element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')  # 设置中文字体

def set_text_formate(doc, text):
    text_para = doc.add_paragraph()
    text_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
    text_para.paragraph_format.line_spacing = Pt(30)  # 行间距，固定值30磅
    text_para.paragraph_format.first_line_indent = Pt(28)  # 首行缩进28磅即空两格
    text_para.paragraph_format.space_before = Pt(0)  # 段前0磅
    text_para.paragraph_format.space_after = Pt(0)  # 段后0磅
    text_formate = text_para.add_run(text)
    text_formate.font.size = Pt(12)  # 设置字体大小-小四号
    text_formate.bold = False  # 设置字体是否加粗
    text_formate.font.name = 'Times New Roman'  # 设置西文字体
    text_formate.element.rPr.rFonts.set(qn('w:eastAsia'), 'GB2312')  # 设置中文字体

if __name__ == '__main__':
    sys.stderr = open('error_log.txt', 'w')

    with open('config.yaml', 'r', encoding="utf-8") as f:  # 用with读取文件更好
        configs = yaml.load(f, Loader=yaml.FullLoader)  # 按字典格式读取并返回
        # 显示读取后的内容
    file_storage_path = configs["file_storage_path"]
    store_text_to_one_word_flag = configs["store_text_to_one_word_flag"]
    waiting_crawl_url = configs["waiting_crawl_url"]
    waiting_crawl_url_start = configs["waiting_crawl_url_start"]
    waiting_crawl_url_end = configs["waiting_crawl_url_end"]
    user_agent = configs["user_agent"]
    get_info_urls_regular_express = configs["get_info_urls_regular_express"]
    get_attachment_name_regular_express = configs["get_attachment_name_regular_express"]
    summarized_data = configs["summarized_data"]
    """伪装正常访问"""
    user = {
        # 打开浏览器输入：“about:version”，查看“用户代理”，即本机的user-agent
        'user-agent': user_agent
    }
    """设置三个全局变量汇总数据"""
    all_title = []
    all_text = []
    all_file_name = []
    first_file_date = ""
    crawl_logs = []
    now_crwal_date_seconds = ""
    now_crwal_date_day = ""
    exited_crwal_date_seconds = ""
    exited_crwal_date_day = ""
    have_new_content_flag = True  # TRUE为本次有新内容需要爬取
    """用range（）函数和for语句来拼接网址"""
    for num in range(waiting_crawl_url_start, waiting_crawl_url_end):  # 用for语句和range（）函数来一次获得数字0-14
        page_url = waiting_crawl_url + str(num)  # 将获得的数字和网址进行拼接得到15个网址，注意要将变量a用str（）转换为字符串才能拼接
        page_response = requests.get(page_url, headers=user)  # 用requests库的get函数访问网页，用headers进行伪装
        page_html = page_response.content.decode('utf-8')  # 用文本显示访问网页得到的内容
        # print(page_html)  # 打印网页内容，以便于使用正则表达式.

        """用正则表达式来获得每个文件的网址"""
        all_info_original_urls = re.findall(get_info_urls_regular_express, page_html)  # 用正则表达式获得文件网址，每个网站可能不同，该参数需要适时修改
        all_info_original_urls_str = "".join(all_info_original_urls)  # 将列表转化为字符串
        all_info_real_urls = re.findall('[a-zA-z]+://[^\s"]*', all_info_original_urls_str)  # 提取网页中的每个文件具体网址

        """用for语句来实现依次访问每个文件的网址，获取正文和标题汇总数据"""
        for index, each_info_url in enumerate(all_info_real_urls):
            response_each_info = requests.get(each_info_url, headers=user)  # 用变量response_gwy保存访问网址后获得的信息
            response_each_info_content = response_each_info.content.decode('utf-8')  # 用'utf-8'的编码模式来记录网址内容，防止出现中文乱码
            each_info_soup = BeautifulSoup(response_each_info_content, features="lxml")
            """获得文件的发布时间"""
            date_day = each_info_soup.find(attrs={"name": "pubDate"})['content'].split(' ')[0]
            date_seconds = each_info_soup.find(attrs={"name": "pubDate"})['content']
            if index == 0:
                first_file_date = date_day
                now_crwal_date_seconds = date_seconds
                now_crwal_date_day = date_day
            """获得文件的标题"""
            # 事例：<meta name="ArticleTitle" content="2023年度山东省公安机关面向社会招录公务员（人民警察）面试成绩（4月20日上午）">
            title = each_info_soup.find(attrs={"name": "ArticleTitle"})['content']  # 用正则表达式来获得每个网址的标题，即文件标题
            file_name = date_day + " " + title

            """获得文件的正文"""
            text_contents = each_info_soup.find("div", id="content_文章").find("div", id="text").contents
            text_list = []
            for text_content in text_contents:
                if type(text_content) == bs4.element.Tag:
                    text_list.append(text_content.contents)

            # 如有之前已经爬取过，则判断本次是否还需要继续爬取,
            if os.path.exists('crawl_log.txt'):
                with open('crawl_log.txt', 'r') as file:
                    existing_data = file.read()
                exited_crwal_date_seconds = existing_data.split("\n")[0].split("#")[0]
                exited_crwal_date_day = exited_crwal_date_seconds.split(" ")[0]
                if date_seconds <= exited_crwal_date_seconds:
                    have_new_content_flag = False  # 本次无新内容需要爬取
                    break
            """标题、正文汇总起来"""
            all_title.append(title)
            all_text.append(text_list)
            all_file_name.append(file_name)
            """爬取记录总起来"""
            crawl_logs.append(date_seconds + "#" + title)

            """获得文件的附件"""
            attachment_urls = each_info_soup.find_all(attrs={"name": "image"})
            if len(attachment_urls) != 0 and store_text_to_one_word_flag == "False":
                # 如果有附件，则新建一个文件夹存放正文和附件
                for num, attachment_url in enumerate(attachment_urls):
                    attachment_real_url_str = attachment_url['content']
                    attachment_type = attachment_real_url_str.split("&")[1].split(".")[1]
                    if not os.path.exists(file_storage_path + file_name):
                        os.makedirs(file_storage_path + file_name)
                    get_each_attachment_name_regular_express = attachment_real_url_str.split("&")[
                                                                   1] + get_attachment_name_regular_express
                    attachment_name_list = re.findall(get_each_attachment_name_regular_express,
                                                      response_each_info_content)

                    attachment_name_str_temp = "".join(attachment_name_list)
                    # 去除<>之间的字符
                    if ">" in attachment_name_str_temp:
                        attachment_name_str = re.sub(r'<[^>]+>', '', attachment_name_str_temp)
                    else:
                        attachment_name_str = attachment_name_str_temp

                    if "附件" in attachment_name_str:
                        attachment_storage_path_all = file_storage_path + file_name + "/" + attachment_name_str
                    else:
                        attachment_storage_path_all = file_storage_path + file_name + "/附件：" + attachment_name_str + "." + attachment_type

                    if "." in attachment_name_str:
                        attachment_storage_path_all = file_storage_path + file_name + "/" + attachment_name_str
                    else:
                        attachment_storage_path_all = file_storage_path + file_name + "/" + attachment_name_str + "." + attachment_type

                    urlretrieve(attachment_real_url_str, attachment_storage_path_all)  # 保存附件

    if have_new_content_flag:  # 只有当有新内容时才会保存
        """将爬取记录保存到txt"""
        if os.path.exists('crawl_log.txt'):
            with open('crawl_log.txt', 'r') as file:
                existing_data = file.read()
        else:
            existing_data = ""
        with open('crawl_log.txt', 'w') as file:
            for log in crawl_logs:
                file.write(log + '\n')
            file.write(existing_data)

        """保存到word"""
        # 方式一：合并到一个word
        if store_text_to_one_word_flag == 'True':
            existed_doc_path_all = file_storage_path + "最新更新时间为：" + exited_crwal_date_day + " " + summarized_data + ".docx"
            new_doc_path_all = file_storage_path + "最新更新时间为：" + now_crwal_date_day + " " + summarized_data + ".docx"
            new_doc = docx.Document()
            for i in range(len(all_file_name)):
                set_title_formate(new_doc, all_title[i])
                for j in range(len(all_text[i])):
                    text_str_temp = ''.join(str(item) for item in all_text[i][j])
                    text_str = re.sub(r'<[^>]+>', '', text_str_temp)
                    set_text_formate(new_doc, text_str)
            new_doc.add_section()  # 分結符
            # 如有不是第一次爬取，则合并两个word
            if not os.path.exists(file_storage_path):
                os.makedirs(file_storage_path)
            if os.path.exists(existed_doc_path_all):
                old_doc = docx.Document(existed_doc_path_all)
                composer = Composer(new_doc)
                composer.append(old_doc)
                composer.save(new_doc_path_all)
            else:
                new_doc.save(new_doc_path_all)
        # 方式二：分别存储到不同word
        else:
            for i, each_file_name in enumerate(all_file_name):
                doc = docx.Document()
                set_title_formate(doc, all_title[i])
                for j in range(len(all_text[i])):
                    text_str_temp = ''.join(str(item) for item in all_text[i][j])
                    text_str = re.sub(r'<[^>]+>', '', text_str_temp)
                    set_text_formate(doc, text_str)
                if os.path.exists(file_storage_path + each_file_name):
                    doc.save(file_storage_path + each_file_name + "/" + each_file_name + ".docx")
                else:
                    doc.save(file_storage_path + each_file_name + ".docx")
