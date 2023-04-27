"""
------------------------------------
# @FileName    :test3.py
# @Time        :2023/4/22 16:19
# @Author      :jack
# @description :
------------------------------------
"""
"""
------------------------------------
# @FileName    :crawl_web_data.py
# @Time        :2023/4/17 19:40
# @Author      :jack
# @description :
------------------------------------
"""
import os
import urllib
from urllib.request import urlretrieve

from docx.oxml.ns import qn
import docx
import requests
from bs4 import BeautifulSoup
from docx import Document  # 安装python-docx
import yaml  # 安装pyyaml
import re  # 导入正则表达式库

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from htmldate import find_date


def set_title_formate(doc, title):
    head = doc.add_heading("", level=1)  # 这里不填标题内容
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 设置为居中
    head_formate = head.add_run(title)
    head_formate.font.name = 'Times New Roman'  # 设置西文字体
    head_formate.element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')  # 设置中文字体


def set_text_formate(doc, cleared_tznr):
    doc.add_paragraph()
    text_para = doc.add_paragraph()
    text_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
    text_para.paragraph_format.line_spacing = Pt(30)  # 行间距，固定值30磅
    text_para.paragraph_format.first_line_indent = Pt(28)  # 首行缩进28磅即空两格
    text_para.paragraph_format.space_before = Pt(0)  # 段前0磅
    text_para.paragraph_format.space_after = Pt(0)  # 段后0磅
    text_formate = text_para.add_run(cleared_tznr)
    text_formate.font.size = Pt(12)  # 设置字体大小-小四号
    text_formate.bold = False  # 设置字体是否加粗
    text_formate.font.name = 'Times New Roman'  # 设置西文字体
    text_formate.element.rPr.rFonts.set(qn('w:eastAsia'), 'GB2312')  # 设置中文字体


if __name__ == '__main__':
    with open('config.yaml', 'r', encoding="utf-8") as f:  # 用with读取文件更好
        configs = yaml.load(f, Loader=yaml.FullLoader)  # 按字典格式读取并返回
        # 显示读取后的内容
    file_storage_path = configs["file_storage_path"]
    store_text_to_one_word_flag = configs["store_text_to_one_word_flag"]
    crawl_text_and_attachment_flag = configs["crawl_text_and_attachment_flag"]
    waiting_crawl_url = configs["waiting_crawl_url"]
    waiting_crawl_url_start = configs["waiting_crawl_url_start"]
    waiting_crawl_url_end = configs["waiting_crawl_url_end"]
    user_agent = configs["user_agent"]
    get_info_urls_regular_express = configs["get_info_urls_regular_express"]
    get_title_regular_express = configs["get_title_regular_express"]
    get_text_regular_express = configs["get_text_regular_express"]
    get_attachment_regular_express = configs["get_attachment_regular_express"]
    get_attachment_name_regular_express = configs["get_attachment_name_regular_express"]
    summarized_data = configs["summarized_data"]
    """伪装正常访问"""
    user = {
        # 打开浏览器输入：“about:version”，查看“用户代理”，即本机的user-agent
        'user-agent': user_agent
    }
    """设置三个全局变量汇总数据"""
    all_title = []
    all_text = []
    all_file_name = []
    first_file_date = ""
    """用range（）函数和for语句来拼接网址"""
    for num in range(waiting_crawl_url_start, waiting_crawl_url_end):  # 用for语句和range（）函数来一次获得数字0-14
        page_url = waiting_crawl_url + str(num)  # 将获得的数字和网址进行拼接得到15个网址，注意要将变量a用str（）转换为字符串才能拼接
        page_response = requests.get(page_url, headers=user)  # 用requests库的get函数访问网页，用headers进行伪装
        page_html = page_response.content.decode('utf-8')  # 用文本显示访问网页得到的内容
        print(page_html)  # 打印网页内容，以便于使用正则表达式.

        """用正则表达式来获得每个文件的网址"""
        all_info_original_urls = re.findall(get_info_urls_regular_express, page_html)  # 用正则表达式获得文件网址，每个网站可能不同，该参数需要适时修改
        all_info_original_urls_str = "".join(all_info_original_urls)  # 将列表转化为字符串
        all_info_real_urls = re.findall('[a-zA-z]+://[^\s"]*', all_info_original_urls_str)  # 提取网页中的每个文件具体网址

        """用for语句来实现依次访问每个文件的网址，获取正文和标题汇总数据"""
        for each_info_url in all_info_real_urls:

            response_each_info = requests.get(each_info_url, headers=user)  # 用变量response_gwy保存访问网址后获得的信息
            response_each_info_content = response_each_info.content.decode('utf-8')  # 用'utf-8'的编码模式来记录网址内容，防止出现中文乱码

            """获得文件的发布时间"""
            date_second = find_date(response_each_info_content, outputformat='%Y-%m-%d %H:%M:%S')
            date_day = find_date(response_each_info_content)
            # if j==0:#第一篇文章时间
            #     first_file_date=date_day
            #
            """用正则表达式获得文件的正文"""
            original_text = re.findall(get_text_regular_express, response_each_info_content)  # 用正则表达式来获得文件的文字内容，清除多余的代码
            original_text_str = ''.join(str(item) for item in original_text).strip()  # 将获得的文字内容由列表转换为字符串
            pattern1 = r'<span .*?>'
            cleared_text = re.sub(pattern1, '', original_text_str)

            """用正则表达式获得文件的标题"""
            # 事例：<meta name="ArticleTitle" content="2023年度山东省公安机关面向社会招录公务员（人民警察）面试成绩（4月20日上午）">
            title = re.findall(get_title_regular_express, response_each_info_content)  # 用正则表达式来获得每个网址的标题，即文件标题
            title = "".join(title)  # 将title列表转换为字符串
            file_name = date_day + title  # 用title字符串来为文件取名字

            """标题、正文汇总起来"""
            all_title.append(title)
            all_text.append(cleared_text)
            all_file_name.append(file_name)
            """用正则表达式获得文件的附件"""
            attachment_original_urls = re.findall(get_attachment_regular_express, response_each_info_content)
            # attachment_name_all_list = re.findall(get_attachment_name_regular_express, response_each_info_content)

            if len(attachment_original_urls) != 0 and store_text_to_one_word_flag=="False":
                attachment_num = 0  # 如果有附件，则新建一个文件夹存放正文和附件
                for attachment_original_url in attachment_original_urls:
                    attachment_num = attachment_num + 1
                    attachment_original_url_str = "".join(attachment_original_url)  # 将列表转化为字符串
                    attachment_real_url = re.findall('[a-zA-z]+://[^\s"]*',
                                                     attachment_original_url_str)  # 提取网页中的每个文件具体网址
                    attachment_real_url_str = "".join(attachment_real_url)
                    attachment_type = attachment_real_url_str.split("&")[1].split(".")[1]
                    if not os.path.exists(file_storage_path + file_name):
                        os.makedirs(file_storage_path + file_name)
                    get_attachment_name_regular_express=attachment_real_url_str.split("&")[1] + get_attachment_name_regular_express
                    attachment_name_list = re.findall(get_attachment_name_regular_express, response_each_info_content)
                    attachment_name_str = "".join(attachment_name_list)
                    attachment_storage_path_all = file_storage_path + file_name + "/" + "附件" + str(
                        attachment_num) + ":" + attachment_name_str + "." + attachment_type
                    urlretrieve(attachment_real_url_str, attachment_storage_path_all)  # 保存附件

    """保存到word"""
    if store_text_to_one_word_flag:
        file_path = "(文章最新更新时间为：)" + first_file_date + file_storage_path + summarized_data + ".docx"

        doc = docx.Document()
        for i in range(len(all_file_name)):
            set_title_formate(doc, all_title[i])
            set_text_formate(doc, all_text[i])
        doc.save(file_path)
    else:
        for i in range(len(all_file_name)):
            doc = docx.Document()
            set_title_formate(doc, all_title[i])
            set_text_formate(doc, all_text[i])
            if os.path.exists(file_storage_path + file_name[i]):
                doc.save(file_storage_path + file_name[i] + "/" + file_name + ".docx")
            else:
                doc.save(file_storage_path + file_name[i] + ".docx")
